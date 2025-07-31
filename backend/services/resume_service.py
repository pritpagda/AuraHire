import io
import re

import PyPDF2
from docx import Document


class ResumeService:
    async def parse_resume(self, file_content: bytes, filename: str) -> str:
        if filename.lower().endswith('.pdf'):
            return await self._parse_pdf(file_content)
        if filename.lower().endswith('.docx'):
            return await self._parse_docx(file_content)
        raise ValueError("Unsupported file format")

    async def _parse_pdf(self, file_content: bytes) -> str:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
        return self._clean_text(text)

    async def _parse_docx(self, file_content: bytes) -> str:
        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)
        text = "".join(p.text + "\n" for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += " ".join(cell.text for cell in row.cells) + "\n"
        return self._clean_text(text)

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        text = text.replace('\x00', '').replace('\r', '\n').strip()
        return text
